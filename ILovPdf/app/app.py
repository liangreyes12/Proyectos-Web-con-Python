import os
import fitz  # PyMuPDF
import tkinter as tk
from tkinter import filedialog, messagebox
from PyPDF2 import PdfReader, PdfWriter
from docx import Document
from fpdf import FPDF

def seleccionar_archivos():
    return filedialog.askopenfilenames(filetypes=[("Archivos PDF", "*.pdf"), ("Archivos Word", "*.docx")])

def guardar_archivo(extension):
    return filedialog.asksaveasfilename(defaultextension=extension, filetypes=[(f"Archivo {extension[1:].upper()}", f"*{extension}")])

def unir_pdfs():
    archivos = seleccionar_archivos()
    if not archivos:
        return
    
    pdf_writer = PdfWriter()
    for archivo in archivos:
        pdf_reader = PdfReader(archivo)
        for pagina in range(len(pdf_reader.pages)):
            pdf_writer.add_page(pdf_reader.pages[pagina])
    
    salida = guardar_archivo(".pdf")
    if salida:
        with open(salida, "wb") as f:
            pdf_writer.write(f)
        messagebox.showinfo("Éxito", "PDFs unidos correctamente.")

def dividir_pdf():
    archivo = filedialog.askopenfilename(filetypes=[("Archivos PDF", "*.pdf")])
    if not archivo:
        return
    
    pdf_reader = PdfReader(archivo)
    num_paginas = len(pdf_reader.pages)
    
    ventana_dividir = tk.Toplevel()
    ventana_dividir.title("Dividir PDF")
    
    tk.Label(ventana_dividir, text=f"El PDF tiene {num_paginas} páginas.").pack()
    
    tk.Label(ventana_dividir, text="Ingrese rango (ej: 1-3) o páginas separadas por comas (ej: 1,3,5)").pack()
    entrada_rango = tk.Entry(ventana_dividir)
    entrada_rango.pack()
    
    def extraer_paginas():
        paginas = entrada_rango.get()
        pdf_writer = PdfWriter()
        
        try:
            if "-" in paginas:
                inicio, fin = map(int, paginas.split("-"))
                paginas_seleccionadas = list(range(inicio - 1, fin))
            else:
                paginas_seleccionadas = [int(p) - 1 for p in paginas.split(",")]
            
            for p in paginas_seleccionadas:
                pdf_writer.add_page(pdf_reader.pages[p])
            
            salida = guardar_archivo(".pdf")
            if salida:
                with open(salida, "wb") as f:
                    pdf_writer.write(f)
                messagebox.showinfo("Éxito", "PDF dividido correctamente.")
        except:
            messagebox.showerror("Error", "Entrada no válida. Use números y comas correctamente.")
    
    tk.Button(ventana_dividir, text="Extraer", command=extraer_paginas).pack()

def pdf_a_word():
    archivo = filedialog.askopenfilename(filetypes=[("Archivos PDF", "*.pdf")])
    if not archivo:
        return
    
    doc = Document()
    pdf = fitz.open(archivo)
    
    for pagina in pdf:
        texto = pagina.get_text("text")
        doc.add_paragraph(texto)
    
    salida = guardar_archivo(".docx")
    if salida:
        doc.save(salida)
        messagebox.showinfo("Éxito", "PDF convertido a Word correctamente.")

def word_a_pdf():
    archivo = filedialog.askopenfilename(filetypes=[("Archivos Word", "*.docx")])
    if not archivo:
        return
    
    doc = Document(archivo)
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    
    for parrafo in doc.paragraphs:
        pdf.cell(200, 10, txt=parrafo.text, ln=True, align='L')
    
    salida = guardar_archivo(".pdf")
    if salida:
        pdf.output(salida)
        messagebox.showinfo("Éxito", "Word convertido a PDF correctamente.")

app = tk.Tk()
app.title("Herramienta PDF/Word")

tk.Button(app, text="Unir PDFs", command=unir_pdfs).pack(fill="x", padx=10, pady=5)
tk.Button(app, text="Dividir PDF", command=dividir_pdf).pack(fill="x", padx=10, pady=5)
tk.Button(app, text="Convertir PDF a Word", command=pdf_a_word).pack(fill="x", padx=10, pady=5)
tk.Button(app, text="Convertir Word a PDF", command=word_a_pdf).pack(fill="x", padx=10, pady=5)

tk.Button(app, text="Salir", command=app.quit).pack(fill="x", padx=10, pady=10)
app.mainloop()